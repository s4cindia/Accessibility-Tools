"""VPAT Editor core — data model + PDF export, transcribed verbatim from the
uploaded vpat_editor.py (GUI-free portions). Single source of truth for the
JBL Cyber Labs report content and the reportlab PDF layout."""
from __future__ import annotations

CONFORMANCE_OPTIONS = [
    "Supports",
    "Partially Supports",
    "Does Not Support",
    "Not Applicable",
    "Not Evaluated",
]

# (bg fill, text colour) used for badges in the UI and cells in the PDF
CONFORMANCE_COLORS = {
    "Supports":          ("#dcfce7", "#166534"),
    "Partially Supports": ("#fef3c7", "#92400e"),
    "Does Not Support":  ("#fee2e2", "#991b1b"),
    "Not Applicable":    ("#e2e8f0", "#475569"),
    "Not Evaluated":     ("#f1f5f9", "#64748b"),
    "":                  ("#f1f5f9", "#64748b"),
}

# Colour used for every hyperlink (criteria links, URLs, email) in UI + PDF
LINK_COLOR = "#1d4ed8"

# Each WCAG success criterion links to its W3C "Understanding" document.
_W3C_22 = "https://www.w3.org/WAI/WCAG22/Understanding/"
_W3C_21 = "https://www.w3.org/WAI/WCAG21/Understanding/"
WCAG_SLUGS = {
    "1.1.1": "non-text-content",
    "1.2.1": "audio-only-and-video-only-prerecorded",
    "1.2.2": "captions-prerecorded",
    "1.2.3": "audio-description-or-media-alternative-prerecorded",
    "1.2.4": "captions-live",
    "1.2.5": "audio-description-prerecorded",
    "1.2.6": "sign-language-prerecorded",
    "1.2.7": "extended-audio-description-prerecorded",
    "1.2.8": "media-alternative-prerecorded",
    "1.2.9": "audio-only-live",
    "1.3.1": "info-and-relationships",
    "1.3.2": "meaningful-sequence",
    "1.3.3": "sensory-characteristics",
    "1.3.4": "orientation",
    "1.3.5": "identify-input-purpose",
    "1.3.6": "identify-purpose",
    "1.4.1": "use-of-color",
    "1.4.2": "audio-control",
    "1.4.3": "contrast-minimum",
    "1.4.4": "resize-text",
    "1.4.5": "images-of-text",
    "1.4.6": "contrast-enhanced",
    "1.4.7": "low-or-no-background-audio",
    "1.4.8": "visual-presentation",
    "1.4.9": "images-of-text-no-exception",
    "1.4.10": "reflow",
    "1.4.11": "non-text-contrast",
    "1.4.12": "text-spacing",
    "1.4.13": "content-on-hover-or-focus",
    "2.1.1": "keyboard",
    "2.1.2": "no-keyboard-trap",
    "2.1.3": "keyboard-no-exception",
    "2.1.4": "character-key-shortcuts",
    "2.2.1": "timing-adjustable",
    "2.2.2": "pause-stop-hide",
    "2.2.3": "no-timing",
    "2.2.4": "interruptions",
    "2.2.5": "re-authenticating",
    "2.2.6": "timeouts",
    "2.3.1": "three-flashes-or-below-threshold",
    "2.3.2": "three-flashes",
    "2.3.3": "animation-from-interactions",
    "2.4.1": "bypass-blocks",
    "2.4.2": "page-titled",
    "2.4.3": "focus-order",
    "2.4.4": "link-purpose-in-context",
    "2.4.5": "multiple-ways",
    "2.4.6": "headings-and-labels",
    "2.4.7": "focus-visible",
    "2.4.8": "location",
    "2.4.9": "link-purpose-link-only",
    "2.4.10": "section-headings",
    "2.4.11": "focus-not-obscured-minimum",
    "2.4.12": "focus-not-obscured-enhanced",
    "2.4.13": "focus-appearance",
    "2.5.1": "pointer-gestures",
    "2.5.2": "pointer-cancellation",
    "2.5.3": "label-in-name",
    "2.5.4": "motion-actuation",
    "2.5.5": "target-size-enhanced",
    "2.5.6": "concurrent-input-mechanisms",
    "2.5.7": "dragging-movements",
    "2.5.8": "target-size-minimum",
    "3.1.1": "language-of-page",
    "3.1.2": "language-of-parts",
    "3.1.3": "unusual-words",
    "3.1.4": "abbreviations",
    "3.1.5": "reading-level",
    "3.1.6": "pronunciation",
    "3.2.1": "on-focus",
    "3.2.2": "on-input",
    "3.2.3": "consistent-navigation",
    "3.2.4": "consistent-identification",
    "3.2.5": "change-on-request",
    "3.2.6": "consistent-help",
    "3.3.1": "error-identification",
    "3.3.2": "labels-or-instructions",
    "3.3.3": "error-suggestion",
    "3.3.4": "error-prevention-legal-financial-data",
    "3.3.5": "help",
    "3.3.6": "error-prevention-all",
    "3.3.7": "redundant-entry",
    "3.3.8": "accessible-authentication-minimum",
    "3.3.9": "accessible-authentication-enhanced",
    "4.1.1": "parsing",          # removed in 2.2 -> use the 2.1 page
    "4.1.2": "name-role-value",
    "4.1.3": "status-messages",
}


def sc_url(criteria_str):
    """Return the W3C Understanding URL for a criterion string, or None."""
    if not criteria_str:
        return None
    sc = criteria_str.strip().split()[0]
    slug = WCAG_SLUGS.get(sc)
    if not slug:
        return None
    base = _W3C_21 if sc == "4.1.1" else _W3C_22
    return base + slug + ".html"

DEFAULT_TERMS = [
    ("Supports",
     "The functionality of the product has at least one method that meets the "
     "criterion without known defects or meets with equivalent facilitation."),
    ("Partially Supports",
     "Some functionality of the product does not meet the criterion."),
    ("Does Not Support",
     "The majority of product functionality does not meet the criterion."),
    ("Not Applicable",
     "The criterion is not relevant to the product."),
    ("Not Evaluated",
     "The product has not been evaluated against the criterion. This can only "
     "be used in WCAG Level AAA criteria."),
]

LEVEL_A = [
    ("1.1.1 Non-text Content (Level A)", "Partially Supports",
     "A few of the alt text are incomplete for the large network diagrams"),
    ("1.2.1 Audio-only and Video-only (Prerecorded) (Level A)", "Supports",
     "Any Audio-only and Video-only context includes textual descriptions that are "
     "available to screen readers or other assistive technology"),
    ("1.2.2 Captions (Prerecorded) (Level A)", "Supports",
     "All video content contains closed captioning"),
    ("1.2.3 Audio Description or Media Alternative (Prerecorded) (Level A)", "Supports",
     "Textual descriptions are made available to screen readers or other assistive technology"),
    ("1.3.1 Info and Relationships (Level A)", "Supports",
     "All visual structures are properly represented within the content"),
    ("1.3.2 Meaningful Sequence (Level A)", "Supports",
     "All tables and visual structures are read in their correct sequence"),
    ("1.3.3 Sensory Characteristics (Level A)", "Supports",
     "Sensory characteristics are not solely used to promote understanding of the content"),
    ("1.4.1 Use of Color (Level A)", "Supports",
     "While color distinctions are used, they are never used as the only visual means "
     "of conveying information"),
    ("1.4.2 Audio Control (Level A)", "Supports",
     "Audio content utilizes proper controls"),
    ("2.1.1 Keyboard (Level A)", "Supports",
     "All content is available and operable through a keyboard interface. Note: User must "
     "turn Accessibility Mode on first"),
    ("2.1.2 No Keyboard Trap (Level A)", "Supports",
     "No keyboard traps exist in the content. Note: User must turn Accessibility Mode on first"),
    ("2.1.4 Character Key Shortcuts (Level A 2.1 and 2.2)", "Supports",
     "Character Key Shortcuts are made available to users"),
    ("2.2.1 Timing Adjustable (Level A)", "Supports",
     "No timing limitations are in place for the content"),
    ("2.2.2 Pause, Stop, Hide (Level A)", "Supports",
     "No timing limitations are in place for this content"),
    ("2.3.1 Three Flashes or Below Threshold (Level A)", "Supports",
     "None of the content utilizes anything that would flash more than three times in any "
     "one second period."),
    ("2.4.1 Bypass Blocks (Level A)", "Partially Supports",
     "Bypass Blocks are missing for the Lab View so users can get immediately to the lab machine"),
    ("2.4.2 Page Titled (Level A)", "Does Not Support",
     "Generic page names are used which make it hard to distinguish between them"),
    ("2.4.3 Focus Order (Level A)", "Partially Supports",
     "All content can be accessed but non-interactive content has also been added to the focus order"),
    ("2.4.4 Link Purpose (In Context) (Level A)", "Supports",
     "The purpose of each link presented can be determined by the link text"),
    ("2.5.1 Pointer Gestures (Level A 2.1 and 2.2)", "Not Applicable",
     "No pointer gestures are utilized within the content"),
    ("2.5.2 Pointer Cancellation (Level A 2.1 and 2.2)", "Not Applicable",
     "No pointer gestures are utilized within the content"),
    ("2.5.3 Label in Name (Level A 2.1 and 2.2)", "Supports",
     "All content is properly labeled. Web content specifically uses properly named elements"),
    ("2.5.4 Motion Actuation (Level A 2.1 and 2.2)", "Not Applicable",
     "No motion actuation is utilized within the content"),
    ("3.1.1 Language of Page (Level A)", "Supports",
     "All content denotes their language as English"),
    ("3.2.1 On Focus (Level A)", "Supports",
     "On Focus indications are made available to users based off the application rendering the "
     "content (i.e. Adobe Reader for a PDF document). Web content specifically uses \u201ctabIndex\u201d "
     "attributes and styling to ensure proper visual indication"),
    ("3.2.2 On Input (Level A)", "Supports",
     "All interactive content controls do not initiate a change of context"),
    ("3.2.6 Consistent Help (Level A 2.2 only)", "Not Evaluated", ""),
    ("3.3.1 Error Identification (Level A)", "Supports",
     "Error identification (when applicable) is made available to users based off the application "
     "rendering the content (i.e. webpage or lab machine)."),
    ("3.3.2 Labels or Instructions (Level A)", "Partially Supports",
     "Instructions on how to/needing to turn on accessibility mode is missing along with a note "
     "that user needs to tab Ctrl button twice to be able to exit the lab machine and get back into "
     "the rest of the page. Additionally, screenshot button is not announcing when it is disabled "
     "but accepts enter/spacebar"),
    ("3.3.7 Redundant Entry (Level A 2.2 only)", "Not Evaluated", ""),
    ("4.1.1 Parsing (Level A)", "Supports",
     "For WCAG 2.0 and 2.1, the September 2023 errata update indicates this criterion is always "
     "supported. See the WCAG 2.0 Editorial Errata and the WCAG 2.1 Editorial Errata."),
    ("4.1.2 Name, Role, Value (Level A)", "Supports",
     "All content employs proper and standard utilization based on its format"),
]

LEVEL_AA = [
    ("1.2.4 Captions (Live) (Level AA)", "Not Applicable", "No live streams are provided"),
    ("1.2.5 Audio Description (Prerecorded) (Level AA)", "Supports",
     "Textual descriptions are made available to screen readers or other assistive technology"),
    ("1.3.4 Orientation (Level AA 2.1 and 2.2)", "Supports", "Orientation works properly"),
    ("1.3.5 Identify Input Purpose (Level AA 2.1 and 2.2)", "Supports",
     "All content contains proper labeling"),
    ("1.4.3 Contrast (Minimum) (Level AA)", "Supports",
     "All images and text meet the ratio of 4.5:1."),
    ("1.4.4 Resize text (Level AA)", "Supports",
     "All content is zoomable up to 200% without pixelization"),
    ("1.4.5 Images of Text (Level AA)", "Supports",
     "Any images of text include an alternative description"),
    ("1.4.10 Reflow (Level AA 2.1 and 2.2)", "Supports",
     "The content does not follow all of the reflow rules due to built-in requirement to show both "
     "the lab guide and lab machine. This is an acceptable exception to the rule."),
    ("1.4.11 Non-text Contrast (Level AA 2.1 and 2.2)", "Supports",
     "All images and text meet the ratio of 4.5:1."),
    ("1.4.12 Text Spacing (Level AA 2.1 and 2.2)", "Supports",
     "All content employs proper text spacing"),
    ("1.4.13 Content on Hover or Focus (Level AA 2.1 and 2.2)", "Supports",
     "No content is offered on Hover or Focus"),
    ("2.4.5 Multiple Ways (Level AA)", "Supports",
     "Access to the lab is made available to users based on a non-linear approach"),
    ("2.4.6 Headings and Labels (Level AA)", "Supports",
     "All visual structures and relationship are properly represented within the content to also be "
     "portrayed to screen readers or other assistive technology."),
    ("2.4.7 Focus Visible (Level AA)", "Supports",
     "Visual focus indicators are made available to users"),
    ("2.4.11 Focus Not Obscured (Minimum) (Level AA 2.2 only)", "Not Evaluated", ""),
    ("2.5.7 Dragging Movements (Level AA 2.2 only)", "Not Evaluated", ""),
    ("2.5.8 Target Size (Minimum) (Level AA 2.2 only)", "Not Evaluated", ""),
    ("3.1.2 Language of Parts (Level AA)", "Supports",
     "All content denotes their language as English"),
    ("3.2.3 Consistent Navigation (Level AA)", "Supports",
     "Access to details of the lab is made available to users via a consistent navigation"),
    ("3.2.4 Consistent Identification (Level AA)", "Supports",
     "Access to details of the lab is made available to users via consistent identification"),
    ("3.3.3 Error Suggestion (Level AA)", "Supports",
     "Error suggestions are not utilized within the content"),
    ("3.3.4 Error Prevention (Legal, Financial, Data) (Level AA)", "Supports",
     "The entry of any legal, financial, or other critical personal data is not utilized within the content"),
    ("3.3.8 Accessible Authentication (Minimum) (Level AA 2.2 only)", "Not Evaluated", ""),
    ("4.1.3 Status Messages (Level AA 2.1 and 2.2)", "Partially Supports",
     "When a lab is started, no announcement is provided to a screen reader"),
]

LEVEL_AAA = [
    ("1.2.6 Sign Language (Prerecorded) (Level AAA)", "Not Evaluated", ""),
    ("1.2.7 Extended Audio Description (Prerecorded) (Level AAA)", "Not Evaluated", ""),
    ("1.2.8 Media Alternative (Prerecorded) (Level AAA)", "Not Evaluated", ""),
    ("1.2.9 Audio-only (Live) (Level AAA)", "Not Evaluated", ""),
    ("1.3.6 Identify Purpose (Level AAA 2.1 and 2.2)", "Not Evaluated", ""),
    ("1.4.6 Contrast (Enhanced) (Level AAA)", "Not Evaluated", ""),
    ("1.4.7 Low or No Background Audio (Level AAA)", "Not Evaluated", ""),
    ("1.4.8 Visual Presentation (Level AAA)", "Not Evaluated", ""),
    ("1.4.9 Images of Text (No Exception) (Level AAA)", "Not Evaluated", ""),
    ("2.1.3 Keyboard (No Exception) (Level AAA)", "Not Evaluated", ""),
    ("2.2.3 No Timing (Level AAA)", "Not Evaluated", ""),
    ("2.2.4 Interruptions (Level AAA)", "Not Evaluated", ""),
    ("2.2.5 Re-authenticating (Level AAA)", "Not Evaluated", ""),
    ("2.2.6 Timeouts (Level AAA 2.1 and 2.2)", "Not Evaluated", ""),
    ("2.3.2 Three Flashes (Level AAA)", "Not Evaluated", ""),
    ("2.3.3 Animation from Interactions (Level AAA 2.1 and 2.2)", "Not Evaluated", ""),
    ("2.4.8 Location (Level AAA)", "Not Evaluated", ""),
    ("2.4.9 Link Purpose (Link Only) (Level AAA)", "Not Evaluated", ""),
    ("2.4.10 Section Headings (Level AAA)", "Not Evaluated", ""),
    ("2.4.12 Focus Not Obscured (Enhanced) (Level AAA 2.2 only)", "Not Evaluated", ""),
    ("2.4.13 Focus Appearance (Level AAA 2.2 only)", "Not Evaluated", ""),
    ("2.5.5 Target Size (Level AAA 2.1 and 2.2)", "Not Evaluated", ""),
    ("2.5.6 Concurrent Input Mechanisms (Level AAA 2.1 and 2.2)", "Not Evaluated", ""),
    ("3.1.3 Unusual Words (Level AAA)", "Not Evaluated", ""),
    ("3.1.4 Abbreviations (Level AAA)", "Not Evaluated", ""),
    ("3.1.5 Reading Level (Level AAA)", "Not Evaluated", ""),
    ("3.1.6 Pronunciation (Level AAA)", "Not Evaluated", ""),
    ("3.2.5 Change on Request (Level AAA)", "Not Evaluated", ""),
    ("3.3.5 Help (Level AAA)", "Not Evaluated", ""),
    ("3.3.6 Error Prevention (All) (Level AAA)", "Not Evaluated", ""),
    ("3.3.9 Accessible Authentication (Enhanced) (Level AAA 2.2 only)", "Not Evaluated", ""),
]

LEGAL_TITLE = "Legal Disclaimer (JBLearning)"
LEGAL_TEXT = (
    "\u00a9 2025. Jones & Bartlett Learning, LLC. All rights reserved. All product and company "
    "names herein are trademarks of their respective owners.\n\n"
    "The information in this document is provided for informational purposes only and is delivered "
    "by Jones & Barlett Learning, LLC (\u201cJBL\u201d) \u201cas is\u201d without warranties of any kind. This "
    "information represents the current view of JBL. Because we are constantly responding to the "
    "market and developing our products and technology, this information is subject to change without "
    "notice and should not be considered a guarantee made by JBL. This is not a contract and no "
    "contractual obligation is formed by this document. While we attempt to keep all information "
    "up-to-date, JBL disclaims any liability with respect to this document or use of the information herein.\n\n"
    "This document contains confidential information of JBL that is only intended for the internal use "
    "of the recipient to whom this was provided by JBL."
)


# Fixed report heading. The title/subtitle are no longer user-editable \u2014 every
# report uses these exact two lines.
REPORT_TITLE = "S4Carlisle Accessibility Conformance Report"
REPORT_SUBTITLE = "WCAG Edition (Based on VPAT\u00ae Version 2.5)"
# The subtitle prints on two centred lines under the title: the edition, then
# the VPAT version on its own line below.
REPORT_SUBTITLE_MAIN = "WCAG Edition"
REPORT_SUBTITLE_VER = "(Based on VPAT\u00ae Version 2.5)"


def default_data():
    """Return a deep-ish copy of the pre-filled Cyber Labs report."""
    return {
        "title": REPORT_TITLE,
        "subtitle": REPORT_SUBTITLE,
        "product": "Cyber Labs",
        "report_date": "January 2025",
        "description": (
            "Cyber Labs consist of a lab guide and the ability to interact with a \u201creal machine\u201d "
            "across the web to be able to train users on cyber security tactics. As such, some of the "
            "material is provided via a hosted website and other pieces of material are provided by "
            "third-party software. This VPAT covers the accessibility of all presented material."
        ),
        "contact": "paul.huntington@ascendlearning.com",
        "notes": (
            "Cyber Labs include interaction with actual instances of Windows and Kali Linux operating "
            "systems. As a third-party company, Microsoft maintains their own accessibility which can be "
            "accessed at: https://support.microsoft.com/en-us/windows/discover-windows-accessibility-"
            "features-8b1068e6-d3b8-4ba8-b027-133dd8911df9.\n\n"
            "Additionally, Kali Linux provides some accessibility features (like a screen reader) that "
            "have been enabled as part of the \u201creal machine\u201d available to the user.\n\n"
            "Each course offers different materials and course material is regularly updated. As such, "
            "should any material be found that does not adhere to the standards expressed herein, please "
            "notify either the above contact or the helpdesk to get it addressed."
        ),
        "eval_methods": (
            "WAVE Accessibility Evaluation Extension (https://wave.webaim.org/extension/)\n"
            "NVDA Screen Reader (https://www.nvaccess.org/download/)\n"
            "Color Contrast Checker (https://webaim.org/resources/contrastchecker/)\n"
            "Adobe Reader\u2019s Read Out Loud (https://www.adobe.com/acrobat/hub/how-to-read-pdf-aloud.html)\n"
            "Microsoft\u2019s Accessibility Checker "
            "(https://support.microsoft.com/en-us/office/improve-accessibility-with-the-accessibility-checker)"
        ),
        "standards": {
            "2.0": {"A": "Yes", "AA": "Yes", "AAA": "No"},
            "2.1": {"A": "Yes", "AA": "Yes", "AAA": "No"},
            "2.2": {"A": "No", "AA": "No", "AAA": "No"},
        },
        "terms": [list(t) for t in DEFAULT_TERMS],
        "level_a": [list(r) for r in LEVEL_A],
        "level_aa": [list(r) for r in LEVEL_AA],
        "level_aaa": [list(r) for r in LEVEL_AAA],
        "legal_title": LEGAL_TITLE,
        "legal_text": LEGAL_TEXT,
    }


# ===========================================================================
# PDF EXPORT (reportlab) — no GUI needed, importable & testable on its own
# ===========================================================================

def export_pdf(data, path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether,
        PageBreak,
    )
    from xml.sax.saxutils import escape as _esc
    import re as _re

    def esc(txt):
        return _esc(str(txt or "")).replace("\n", "<br/>")

    _URL_RE = _re.compile(r'(https?://[^\s<>"\')]+)')
    _EMAIL_RE = _re.compile(r'([A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})')

    def _link(href, text):
        return '<a href="%s" color="%s"><u>%s</u></a>' % (href, LINK_COLOR, _esc(text))

    def linkify(txt):
        """Escape text and turn http(s) URLs and emails into blue links."""
        txt = str(txt or "")
        out, last = [], 0
        # combined scan for URLs first, then emails inside the gaps
        for m in _URL_RE.finditer(txt):
            out.append(_emailify(txt[last:m.start()]))
            url = m.group(1)
            trail = ""
            while url and url[-1] in ".,;:)]":
                trail = url[-1] + trail
                url = url[:-1]
            out.append(_link(url, url) + _esc(trail))
            last = m.end()
        out.append(_emailify(txt[last:]))
        return "".join(out).replace("\n", "<br/>")

    def _emailify(txt):
        out, last = [], 0
        for m in _EMAIL_RE.finditer(txt):
            out.append(_esc(txt[last:m.start()]))
            addr = m.group(1)
            out.append(_link("mailto:" + addr, addr))
            last = m.end()
        out.append(_esc(txt[last:]))
        return "".join(out)

    NAVY = colors.HexColor("#1e3a5f")
    INK = colors.HexColor("#1f2937")
    GRID = colors.HexColor("#cbd5e1")
    HEADBG = colors.HexColor("#1e3a5f")

    styles = getSampleStyleSheet()
    s_title = ParagraphStyle("vt", parent=styles["Title"], fontName="Helvetica-Bold",
                             fontSize=18, textColor=NAVY, alignment=TA_CENTER, spaceAfter=2)
    s_sub = ParagraphStyle("vs", parent=styles["Normal"], fontName="Helvetica",
                           fontSize=11, textColor=colors.HexColor("#475569"),
                           alignment=TA_CENTER, spaceAfter=14)
    s_sub_main = ParagraphStyle("vs1", parent=styles["Normal"], fontName="Helvetica-Bold",
                                fontSize=13, textColor=colors.HexColor("#334155"),
                                alignment=TA_CENTER, spaceAfter=0)
    s_sub_ver = ParagraphStyle("vs2", parent=styles["Normal"], fontName="Helvetica",
                               fontSize=10.5, textColor=colors.HexColor("#475569"),
                               alignment=TA_CENTER, spaceAfter=0)
    s_h2 = ParagraphStyle("vh2", parent=styles["Heading2"], fontName="Helvetica-Bold",
                          fontSize=13, textColor=NAVY, spaceBefore=14, spaceAfter=6)
    s_label = ParagraphStyle("vlbl", parent=styles["Normal"], fontName="Helvetica-Bold",
                             fontSize=12, textColor=INK, spaceAfter=2)
    s_body = ParagraphStyle("vbody", parent=styles["Normal"], fontName="Helvetica",
                            fontSize=12, textColor=INK, leading=16, spaceAfter=8)
    s_cell = ParagraphStyle("vcell", parent=styles["Normal"], fontName="Helvetica",
                            fontSize=8.5, textColor=INK, leading=11)
    s_cellb = ParagraphStyle("vcellb", parent=styles["Normal"], fontName="Helvetica-Bold",
                             fontSize=8.5, textColor=INK, leading=11)
    s_hdr = ParagraphStyle("vhdr", parent=styles["Normal"], fontName="Helvetica-Bold",
                           fontSize=9, textColor=colors.white, leading=11)
    s_small = ParagraphStyle("vsm", parent=styles["Normal"], fontName="Helvetica",
                             fontSize=7.5, textColor=colors.HexColor("#64748b"),
                             alignment=TA_CENTER, leading=10)

    story = []

    # ---- title block (fixed heading). The ITI trademark note now lives in the
    # page footer, directly above the page number (see `footer` below). The
    # subtitle prints on two centred lines: "WCAG Edition" then, on the line
    # below, "(Based on VPAT® Version 2.5)". ----
    story.append(Paragraph(esc(REPORT_TITLE), s_title))
    story.append(Spacer(1, 6))
    story.append(Paragraph(esc(REPORT_SUBTITLE_MAIN), s_sub_main))
    story.append(Spacer(1, 13))                  # one blank line
    story.append(Paragraph(esc(REPORT_SUBTITLE_VER), s_sub_ver))
    story.append(Spacer(1, 14))

    def field(label, value, link=False):
        story.append(Paragraph(esc(label), s_label))
        story.append(Paragraph(linkify(value) if link else esc(value), s_body))

    # "Name of Product/Version: <value>" on one line, then a blank line, then
    # "Report Date: <value>" on its own line (label and value inline in both).
    story.append(Paragraph(
        f'<b>Name of Product/Version:</b> {esc(data["product"])}', s_body))
    story.append(Spacer(1, 6))                   # one blank line
    story.append(Paragraph(
        f'<b>Report Date:</b> {esc(data["report_date"])}', s_body))
    field("Product Description:", data["description"], link=True)
    # "Contact Information: <value>" on a single line (label and value inline).
    story.append(Paragraph(
        f'<b>Contact Information:</b> {linkify(data["contact"])}', s_body))
    field("Notes:", data["notes"], link=True)

    story.append(Paragraph("Evaluation Methods Used:", s_label))
    methods = [m.strip() for m in data["eval_methods"].splitlines() if m.strip()]
    bullets = "<br/>".join("\u2022 " + linkify(m) for m in methods)
    story.append(Paragraph(bullets, s_body))

    # ---- applicable standards ----
    story.append(Paragraph("Applicable Standards/Guidelines", s_h2))
    story.append(Paragraph(
        "This report covers the degree of conformance for the following accessibility "
        "standard/guidelines:", s_body))
    std_rows = [[Paragraph("Standard/Guideline", s_hdr), Paragraph("Included In Report", s_hdr)]]
    ver_names = {"2.0": "Web Content Accessibility Guidelines 2.0",
                 "2.1": "Web Content Accessibility Guidelines 2.1",
                 "2.2": "Web Content Accessibility Guidelines 2.2"}
    for ver in ["2.0", "2.1", "2.2"]:
        st = data["standards"][ver]
        inc = (f"Level A ({st['A']})<br/>Level AA ({st['AA']})<br/>Level AAA ({st['AAA']})")
        std_rows.append([Paragraph(ver_names[ver], s_cellb), Paragraph(inc, s_cell)])
    std_tbl = Table(std_rows, colWidths=[3.4 * inch, 3.4 * inch])
    std_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HEADBG),
        ("GRID", (0, 0), (-1, -1), 0.5, GRID),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
    ]))
    story.append(std_tbl)

    # ---- terms ----
    story.append(Paragraph("Terms", s_h2))
    story.append(Paragraph(
        "The terms used in the Conformance Level information are defined as follows:", s_body))
    term_bits = []
    for term, definition in data["terms"]:
        term_bits.append("\u2022 <b>%s:</b> %s" % (_esc(term), _esc(definition)))
    story.append(Paragraph("<br/>".join(term_bits), s_body))

    story.append(Paragraph(
        "Note: When reporting on conformance with the WCAG 2.x Success Criteria, they are scoped for "
        "full pages, complete processes, and accessibility-supported ways of using technology as "
        "documented in the WCAG 2.0 Conformance Requirements.", s_body))

    # ---- success-criteria tables ----
    def build_criteria_table(title, rows):
        elems = [Paragraph(title, s_h2)]
        tbl_rows = [[Paragraph("Criteria", s_hdr),
                     Paragraph("Conformance Level", s_hdr),
                     Paragraph("Remarks and Explanations", s_hdr)]]
        styling = [
            ("BACKGROUND", (0, 0), (-1, 0), HEADBG),
            ("GRID", (0, 0), (-1, -1), 0.5, GRID),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ]
        for i, (crit, conf, remarks) in enumerate(rows, start=1):
            bg_hex, _ = CONFORMANCE_COLORS.get(conf, CONFORMANCE_COLORS[""])
            url = sc_url(crit)
            crit_para = _link(url, crit) if url else esc(crit)
            tbl_rows.append([
                Paragraph(crit_para, s_cellb),
                Paragraph(esc(conf), s_cell),
                Paragraph(esc(remarks), s_cell),
            ])
            styling.append(("BACKGROUND", (1, i), (1, i), colors.HexColor(bg_hex)))
        tbl = Table(tbl_rows, colWidths=[2.45 * inch, 1.35 * inch, 3.0 * inch], repeatRows=1)
        tbl.setStyle(TableStyle(styling))
        elems.append(tbl)
        return elems

    # Table 1 always begins on a fresh page (keeps the overview/terms section
    # and the success-criteria tables cleanly separated).
    story.append(PageBreak())
    story += build_criteria_table("Table 1: Success Criteria, Level A", data["level_a"])
    story += build_criteria_table("Table 2: Success Criteria, Level AA", data["level_aa"])
    story += build_criteria_table("Table 3: Success Criteria, Level AAA", data["level_aaa"])

    # ---- legal ----
    story.append(Spacer(1, 8))
    story.append(KeepTogether([
        Paragraph(esc(data["legal_title"]), s_h2),
        Paragraph(esc(data["legal_text"]), s_body),
    ]))

    _TRADEMARK = ("“Voluntary Product Accessibility Template” and “VPAT” are "
                  "registered service marks of the Information Technology "
                  "Industry Council (ITI)")

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#94a3b8"))
        # ITI trademark note — only on the FIRST page, just above the page
        # number (shrunk to fit the text width).
        if doc.page == 1:
            tsize = 7.0
            usable = letter[0] - 1.4 * inch      # page width minus L/R margins
            while tsize > 5 and canvas.stringWidth(_TRADEMARK, "Helvetica", tsize) > usable:
                tsize -= 0.5
            canvas.setFont("Helvetica", tsize)
            canvas.drawCentredString(letter[0] / 2.0, 0.60 * inch, _TRADEMARK)
        # Page number on every page.
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(letter[0] / 2.0, 0.40 * inch, "Page %d" % doc.page)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        path, pagesize=letter,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.6 * inch, bottomMargin=0.85 * inch,
        title="%s - %s" % (REPORT_TITLE, data["product"]),
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return path


# ===========================================================================
# WORD EXPORT (python-docx) — mirrors the PDF layout so users can download the
# finished report as an editable .docx as well as a PDF.
# ===========================================================================

def export_docx(data, path):
    import re as _re
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    _URL_RE = _re.compile(r'(https?://[^\s<>"\')]+)')
    _EMAIL_RE = _re.compile(r'([A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})')

    NAVY = RGBColor(0x1E, 0x3A, 0x5F)
    INK = RGBColor(0x1F, 0x29, 0x37)
    LINK_RGB = RGBColor(0x1D, 0x4E, 0xD8)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    def _shade(cell, hex_fill):
        """Set a table cell's background colour (hex without '#')."""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        tcPr.append(shd)

    def _add_hyperlink(paragraph, url, text):
        """Append a clickable hyperlink run to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        hyper = OxmlElement("w:hyperlink")
        hyper.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1D4ED8")
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        new_run.append(t)
        hyper.append(new_run)
        paragraph._p.append(hyper)

    def _add_linked_text(paragraph, text, size=12, bold=False, color=INK):
        """Add text to a paragraph, turning URLs and emails into hyperlinks.
        Newlines become line breaks within the paragraph."""
        text = str(text or "")
        # tokenise into (kind, value): url / email / plain
        tokens = []
        last = 0
        for m in _URL_RE.finditer(text):
            if m.start() > last:
                tokens.append(("plain", text[last:m.start()]))
            url = m.group(1)
            trail = ""
            while url and url[-1] in ".,;:)]":
                trail = url[-1] + trail
                url = url[:-1]
            tokens.append(("url", url))
            if trail:
                tokens.append(("plain", trail))
            last = m.end()
        tokens.append(("plain", text[last:]))

        def _emit_plain(chunk):
            # inside plain chunks, still linkify emails
            elast = 0
            for em in _EMAIL_RE.finditer(chunk):
                if em.start() > elast:
                    _emit_run(chunk[elast:em.start()])
                _add_hyperlink(paragraph, "mailto:" + em.group(1), em.group(1))
                elast = em.end()
            if elast < len(chunk):
                _emit_run(chunk[elast:])

        def _emit_run(chunk):
            parts = chunk.split("\n")
            for i, seg in enumerate(parts):
                if i:
                    paragraph.add_run().add_break()
                if seg:
                    run = paragraph.add_run(seg)
                    run.font.size = Pt(size)
                    run.font.bold = bold
                    run.font.color.rgb = color

        for kind, val in tokens:
            if kind == "url":
                _add_hyperlink(paragraph, val, val)
            else:
                _emit_plain(val)

    def _para(text="", size=12, bold=False, color=INK, align=None,
              space_after=8):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color
        return p

    def _heading(text):
        p = _para(text, size=13, bold=True, color=NAVY, space_after=6)
        p.paragraph_format.space_before = Pt(14)
        return p

    def _label_value(label, value, link=False):
        """'Label: value' with the label bold and the value inline."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(label + " ")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = INK
        if link:
            _add_linked_text(p, value, size=12)
        else:
            r = p.add_run(str(value or ""))
            r.font.size = Pt(12)
            r.font.color.rgb = INK
        return p

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.85)

    # ---- title block ----
    _para(REPORT_TITLE, size=18, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _para(REPORT_SUBTITLE_MAIN, size=13, bold=True,
          color=RGBColor(0x33, 0x41, 0x55),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _para(REPORT_SUBTITLE_VER, size=10.5, color=RGBColor(0x47, 0x55, 0x69),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    # ---- overview fields (labels + values inline where the PDF is inline) ----
    _label_value("Name of Product/Version:", data["product"])
    _label_value("Report Date:", data["report_date"])
    _para("Product Description:", size=12, bold=True, space_after=2)
    _add_linked_text(_para(space_after=8), data["description"])
    # "Contact Information: <value>" on a single line.
    _label_value("Contact Information:", data["contact"], link=True)
    _para("Notes:", size=12, bold=True, space_after=2)
    _add_linked_text(_para(space_after=8), data["notes"])

    _para("Evaluation Methods Used:", size=12, bold=True, space_after=2)
    methods = [m.strip() for m in data["eval_methods"].splitlines() if m.strip()]
    for m in methods:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        _add_linked_text(bp, m, size=12)

    # ---- applicable standards ----
    _heading("Applicable Standards/Guidelines")
    _para("This report covers the degree of conformance for the following "
          "accessibility standard/guidelines:", size=12)
    ver_names = {"2.0": "Web Content Accessibility Guidelines 2.0",
                 "2.1": "Web Content Accessibility Guidelines 2.1",
                 "2.2": "Web Content Accessibility Guidelines 2.2"}
    std_tbl = doc.add_table(rows=1, cols=2)
    std_tbl.style = "Table Grid"
    std_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = std_tbl.rows[0].cells
    for cell, label in zip(hdr, ("Standard/Guideline", "Included In Report")):
        _shade(cell, "1E3A5F")
        cp = cell.paragraphs[0]
        run = cp.add_run(label)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
    for ver in ["2.0", "2.1", "2.2"]:
        st = data["standards"][ver]
        row = std_tbl.add_row().cells
        r0 = row[0].paragraphs[0].add_run(ver_names[ver])
        r0.font.bold = True
        r0.font.size = Pt(8.5)
        inc = f"Level A ({st['A']})\nLevel AA ({st['AA']})\nLevel AAA ({st['AAA']})"
        cp = row[1].paragraphs[0]
        for i, line in enumerate(inc.split("\n")):
            if i:
                cp.add_run().add_break()
            cp.add_run(line).font.size = Pt(8.5)

    # ---- terms ----
    _heading("Terms")
    _para("The terms used in the Conformance Level information are defined as "
          "follows:", size=12)
    for term, definition in data["terms"]:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        r = bp.add_run(term + ": ")
        r.font.bold = True
        r.font.size = Pt(12)
        bp.add_run(definition).font.size = Pt(12)

    _para("Note: When reporting on conformance with the WCAG 2.x Success "
          "Criteria, they are scoped for full pages, complete processes, and "
          "accessibility-supported ways of using technology as documented in "
          "the WCAG 2.0 Conformance Requirements.", size=12)

    # ---- success-criteria tables ----
    def build_criteria_table(title, rows):
        _heading(title)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        widths = (Inches(2.45), Inches(1.35), Inches(3.0))
        hdr = tbl.rows[0].cells
        for cell, label, w in zip(
                hdr, ("Criteria", "Conformance Level",
                      "Remarks and Explanations"), widths):
            cell.width = w
            _shade(cell, "1E3A5F")
            run = cell.paragraphs[0].add_run(label)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
        for crit, conf, remarks in rows:
            cells = tbl.add_row().cells
            for cell, w in zip(cells, widths):
                cell.width = w
            # criteria (linked to W3C when known)
            cp = cells[0].paragraphs[0]
            url = sc_url(crit)
            if url:
                _add_hyperlink(cp, url, str(crit))
            else:
                r = cp.add_run(str(crit))
                r.font.bold = True
                r.font.size = Pt(8.5)
            # conformance level, shaded by status
            bg_hex, _ = CONFORMANCE_COLORS.get(conf, CONFORMANCE_COLORS[""])
            _shade(cells[1], bg_hex.lstrip("#"))
            cells[1].paragraphs[0].add_run(str(conf or "")).font.size = Pt(8.5)
            # remarks (linkified)
            _add_linked_text(cells[2].paragraphs[0], remarks, size=8.5)

    # Table 1 always begins on a fresh page.
    doc.add_page_break()
    build_criteria_table("Table 1: Success Criteria, Level A", data["level_a"])
    build_criteria_table("Table 2: Success Criteria, Level AA", data["level_aa"])
    build_criteria_table("Table 3: Success Criteria, Level AAA", data["level_aaa"])

    # ---- legal ----
    _heading(data["legal_title"])
    _add_linked_text(_para(space_after=8), data["legal_text"])

    doc.save(path)
    return path
